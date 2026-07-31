from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def emu_to_inches(value):
    return round(value / 914400, 4) if value is not None else None


def paragraph_info(index, paragraph):
    ppr = paragraph._p.pPr
    rpr = paragraph.runs[0]._r.rPr if paragraph.runs else None
    fonts = {}
    if rpr is not None and rpr.rFonts is not None:
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            value = rpr.rFonts.get(qn(f"w:{key}"))
            if value:
                fonts[key] = value
    return {
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "alignment": str(paragraph.alignment),
        "format": {
            "left_indent": emu_to_inches(paragraph.paragraph_format.left_indent),
            "right_indent": emu_to_inches(paragraph.paragraph_format.right_indent),
            "first_line_indent": emu_to_inches(
                paragraph.paragraph_format.first_line_indent
            ),
            "space_before_pt": (
                round(paragraph.paragraph_format.space_before.pt, 2)
                if paragraph.paragraph_format.space_before
                else None
            ),
            "space_after_pt": (
                round(paragraph.paragraph_format.space_after.pt, 2)
                if paragraph.paragraph_format.space_after
                else None
            ),
            "line_spacing": str(paragraph.paragraph_format.line_spacing),
            "keep_with_next": paragraph.paragraph_format.keep_with_next,
            "page_break_before": paragraph.paragraph_format.page_break_before,
        },
        "first_run": {
            "bold": paragraph.runs[0].bold if paragraph.runs else None,
            "italic": paragraph.runs[0].italic if paragraph.runs else None,
            "size_pt": (
                round(paragraph.runs[0].font.size.pt, 2)
                if paragraph.runs and paragraph.runs[0].font.size
                else None
            ),
            "font_name": paragraph.runs[0].font.name if paragraph.runs else None,
            "fonts": fonts,
        },
        "has_ppr": ppr is not None,
    }


def main():
    source = Path(sys.argv[1]).resolve()
    output = Path(sys.argv[2]).resolve()
    document = Document(source)

    package_parts = []
    with zipfile.ZipFile(source) as archive:
        for info in sorted(archive.infolist(), key=lambda item: item.filename):
            payload = archive.read(info.filename)
            package_parts.append(
                {
                    "path": info.filename,
                    "size": len(payload),
                    "sha256": hashlib.sha256(payload).hexdigest(),
                }
            )

    tables = []
    for table_index, table in enumerate(document.tables):
        tables.append(
            {
                "index": table_index,
                "style": table.style.name if table.style else None,
                "rows": [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ],
            }
        )

    sections = []
    for index, section in enumerate(document.sections):
        sections.append(
            {
                "index": index,
                "page_width_in": emu_to_inches(section.page_width),
                "page_height_in": emu_to_inches(section.page_height),
                "top_margin_in": emu_to_inches(section.top_margin),
                "right_margin_in": emu_to_inches(section.right_margin),
                "bottom_margin_in": emu_to_inches(section.bottom_margin),
                "left_margin_in": emu_to_inches(section.left_margin),
                "header_distance_in": emu_to_inches(section.header_distance),
                "footer_distance_in": emu_to_inches(section.footer_distance),
                "different_first_page": section.different_first_page_header_footer,
                "start_type": str(section.start_type),
                "header": [p.text for p in section.header.paragraphs],
                "first_page_header": [
                    p.text for p in section.first_page_header.paragraphs
                ],
                "footer": [p.text for p in section.footer.paragraphs],
                "first_page_footer": [
                    p.text for p in section.first_page_footer.paragraphs
                ],
            }
        )

    inventory = {
        "source": str(source),
        "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "paragraph_count": len(document.paragraphs),
        "paragraphs": [
            paragraph_info(index, paragraph)
            for index, paragraph in enumerate(document.paragraphs)
        ],
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(document.inline_shapes),
        "package_parts": package_parts,
    }
    output.write_text(
        json.dumps(inventory, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
