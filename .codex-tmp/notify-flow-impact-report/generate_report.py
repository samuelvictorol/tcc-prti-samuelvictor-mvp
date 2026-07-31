from __future__ import annotations

import hashlib
import os
import re
import shutil
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\samuel\tcc\notify-app")
REFERENCE = Path(
    r"C:\Users\samuel\Downloads\Template Relatorio de Impacto - TJGO PRTI - Samuel Victor Oliveira Lima.docx"
)
WORK = ROOT / ".codex-tmp" / "notify-flow-impact-report"
ASSETS = WORK / "generated-assets"
OUTPUT = ROOT / "docs" / "Relatorio-de-Impacto-Notify-Flow-Gate-4-Samuel-Victor-Oliveira-Lima.docx"
PROFILE_IMAGE = ROOT / "frontend" / "public" / "meuperfil.png"

TITLE = (
    "NOTIFY FLOW: RELATÓRIO DE IMPACTO, TESTES E VALIDAÇÃO DE UMA "
    "PLATAFORMA MULTICANAL DE NOTIFICAÇÕES CONSENTIDAS"
)
SHORT_TITLE = "Notify Flow — Relatório de Impacto, Testes e Validação"
AUTHOR = "Samuel Victor Oliveira Lima"
ADVISOR = "Leonardo Oliveira"

INK = "#082B28"
TEAL = "#1FB7A6"
TEAL_DARK = "#087E73"
MINT = "#E8FAF6"
BLUE = "#2AABEE"
BLUE_LIGHT = "#EAF7FE"
GREEN = "#25D366"
GREEN_DARK = "#128C7E"
GREEN_LIGHT = "#E9FAF0"
RED = "#D94B4B"
RED_LIGHT = "#FFF0F0"
AMBER = "#D99116"
AMBER_LIGHT = "#FFF7E8"
GRAY = "#5F716F"
LIGHT = "#F5FAF9"
LINE = "#CFE3DF"
WHITE = "#FFFFFF"


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str = LINE,
    radius: int = 26,
    width: int = 3,
) -> None:
    draw.rounded_rectangle(
        xy,
        radius=radius,
        fill=hex_rgb(fill),
        outline=hex_rgb(outline),
        width=width,
    )


def text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = INK,
    spacing: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=spacing)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - w) / 2, (y1 + y2 - h) / 2),
        text,
        font=font,
        fill=hex_rgb(fill),
        align="center",
        spacing=spacing,
    )


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = TEAL_DARK,
    width: int = 8,
) -> None:
    draw.line([start, end], fill=hex_rgb(color), width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    mag = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / mag, dy / mag
    px, py = -uy, ux
    length, spread = 24, 13
    p1 = (x2 - ux * length + px * spread, y2 - uy * length + py * spread)
    p2 = (x2 - ux * length - px * spread, y2 - uy * length - py * spread)
    draw.polygon([end, p1, p2], fill=hex_rgb(color))


def canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1000), hex_rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1800, 16), fill=hex_rgb(TEAL))
    draw.text((80, 55), title, font=get_font(48, True), fill=hex_rgb(INK))
    draw.text((80, 120), subtitle, font=get_font(25), fill=hex_rgb(GRAY))
    return image, draw


def save_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out: dict[str, Path] = {}

    # 1. Architecture
    image, draw = canvas(
        "Arquitetura ponta a ponta",
        "Separação em camadas, integrações oficiais, processamento assíncrono e auditoria",
    )
    boxes = {
        "Usuários": (70, 260, 350, 505),
        "Frontend": (465, 215, 780, 550),
        "API": (900, 215, 1215, 550),
        "Provedores": (1330, 180, 1720, 585),
        "Dados": (900, 680, 1215, 900),
        "Tempo real": (465, 680, 780, 900),
    }
    rounded_box(draw, boxes["Usuários"], BLUE_LIGHT, BLUE)
    rounded_box(draw, boxes["Frontend"], MINT, TEAL)
    rounded_box(draw, boxes["API"], LIGHT, TEAL_DARK)
    rounded_box(draw, boxes["Provedores"], GREEN_LIGHT, GREEN_DARK)
    rounded_box(draw, boxes["Dados"], AMBER_LIGHT, AMBER)
    rounded_box(draw, boxes["Tempo real"], RED_LIGHT, RED)
    text_center(draw, boxes["Usuários"], "ADMINISTRADOR\n\nCONTATO", get_font(31, True))
    text_center(
        draw,
        boxes["Frontend"],
        "VUE 3 + QUASAR\n\nRotas públicas\nPainel administrativo\nMeu Perfil",
        get_font(29, True),
    )
    text_center(
        draw,
        boxes["API"],
        "EXPRESS + MONGOOSE\n\nControllers\nManagers\nDTOs e políticas\nServiços",
        get_font(28, True),
    )
    text_center(
        draw,
        boxes["Provedores"],
        "WHATSAPP CLOUD\nWebhook + Graph API\n\nTELEGRAM BOT API\n\nGMAIL SMTP",
        get_font(27, True),
    )
    text_center(
        draw,
        boxes["Dados"],
        "MONGODB\nDados e auditoria\n\nREDIS + BULLMQ\nFila e deduplicação",
        get_font(27, True),
    )
    text_center(
        draw,
        boxes["Tempo real"],
        "SOCKET.IO\n\nMensagens\nRecibos\nAlertas\nLogs",
        get_font(29, True),
    )
    arrow(draw, (350, 380), (465, 380))
    arrow(draw, (780, 380), (900, 380))
    arrow(draw, (1215, 380), (1330, 380))
    arrow(draw, (1058, 550), (1058, 680))
    arrow(draw, (900, 790), (780, 790))
    draw.text(
        (72, 925),
        "HTTPS/JSON • JWT/cookies seguros • Webhooks assinados • dados sensíveis criptografados",
        font=get_font(24, True),
        fill=hex_rgb(TEAL_DARK),
    )
    path = ASSETS / "figura-arquitetura.png"
    image.save(path, quality=95)
    out["architecture"] = path

    # 2. Consent and identity
    image, draw = canvas(
        "Jornada de consentimento e vínculo",
        "O contato inicia a interação; cada canal mantém autorização independente e revogável",
    )
    steps = [
        ("1", "CONVITE OU\nCOMANDO", "Link rastreável,\n/notify-me ou\n/verify-me", BLUE_LIGHT, BLUE),
        ("2", "IDENTIDADE REAL", "Webhook recebe\nchat_id, wa_id\nou telefone", GREEN_LIGHT, GREEN_DARK),
        ("3", "VÍNCULO SEGURO", "Telefone compartilhado,\ne-mail validado\ne deduplicação", MINT, TEAL),
        ("4", "CONSENTIMENTO", "Evento registra canal,\norigem, finalidade\ne versão do termo", AMBER_LIGHT, AMBER),
        ("5", "AUTOGESTÃO", "Meu Perfil:\nrevogar, consultar\ne remover vínculos", RED_LIGHT, RED),
    ]
    x = 55
    box_w = 300
    for number, heading, body, fill, outline in steps:
        rounded_box(draw, (x, 260, x + box_w, 690), fill, outline)
        draw.ellipse((x + 105, 290, x + 195, 380), fill=hex_rgb(outline))
        text_center(draw, (x + 105, 290, x + 195, 380), number, get_font(38, True), WHITE)
        text_center(draw, (x + 25, 405, x + box_w - 25, 500), heading, get_font(27, True))
        text_center(draw, (x + 25, 515, x + box_w - 25, 655), body, get_font(24))
        if number != "5":
            arrow(draw, (x + box_w, 475), (x + box_w + 55, 475), TEAL_DARK, 6)
        x += 355
    rounded_box(draw, (240, 775, 1560, 920), LIGHT, LINE, 18, 2)
    text_center(
        draw,
        (260, 790, 1540, 905),
        "Princípio operacional: não inferir autorização pela simples existência do contato.\n"
        "Telegram, WhatsApp Cloud e e-mail são ajustáveis separadamente e toda revogação fica auditada.",
        get_font(26, True),
        TEAL_DARK,
    )
    path = ASSETS / "figura-consentimento.png"
    image.save(path, quality=95)
    out["consent"] = path

    # 3. Queue lifecycle
    image, draw = canvas(
        "Ciclo de entrega confiável",
        "Persistência, fila, tentativas controladas, recibos do provedor e trilha por destinatário",
    )
    nodes = [
        ((70, 250, 340, 500), "CAMPANHA", "Contatos, grupos,\ntemplates e variáveis", BLUE_LIGHT, BLUE),
        ((440, 250, 710, 500), "MONGODB", "Entrega criada como\npendente de enfileirar", MINT, TEAL),
        ((810, 250, 1080, 500), "BULLMQ", "Redis, concorrência 5,\n4 tentativas", AMBER_LIGHT, AMBER),
        ((1180, 250, 1450, 500), "PROVEDOR", "WhatsApp, Telegram\nou Gmail", GREEN_LIGHT, GREEN_DARK),
        ((1515, 250, 1760, 500), "RECIBO", "Enviado, entregue,\nlido ou falhou", RED_LIGHT, RED),
    ]
    for i, (box, title, body, fill, outline) in enumerate(nodes):
        rounded_box(draw, box, fill, outline)
        text_center(draw, (box[0] + 15, box[1] + 25, box[2] - 15, box[1] + 105), title, get_font(28, True))
        text_center(draw, (box[0] + 20, box[1] + 105, box[2] - 20, box[3] - 25), body, get_font(24))
        if i < len(nodes) - 1:
            arrow(draw, (box[2], 375), (nodes[i + 1][0][0], 375), TEAL_DARK, 6)
    rounded_box(draw, (255, 650, 750, 890), RED_LIGHT, RED)
    text_center(
        draw,
        (280, 675, 725, 860),
        "FALHA TRANSITÓRIA\n\nBackoff exponencial\n2 s → 4 s → 8 s\nsem interromper a campanha",
        get_font(27, True),
    )
    rounded_box(draw, (1050, 650, 1545, 890), LIGHT, TEAL_DARK)
    text_center(
        draw,
        (1075, 675, 1520, 860),
        "AUDITORIA\n\nStatus por contato e canal,\ntentativas, motivo e requestId\ncom segredos redigidos",
        get_font(27, True),
    )
    arrow(draw, (945, 500), (750, 650), RED, 6)
    arrow(draw, (1635, 500), (1545, 650), TEAL_DARK, 6)
    path = ASSETS / "figura-fila.png"
    image.save(path, quality=95)
    out["queue"] = path

    # 4. Meta onboarding
    image, draw = canvas(
        "Preparação do WhatsApp Cloud para produção",
        "A sequência combina requisitos empresariais, segurança, credenciais e homologação de modelos",
    )
    rows = [
        ("1", "BASE EMPRESARIAL", "Empresa e CNPJ reais • conta Facebook • Página • portfólio empresarial", BLUE_LIGHT, BLUE),
        ("2", "VERIFICAÇÃO E SEGURANÇA", "Documentos legíveis • domínio e dados consistentes • autenticação em dois fatores", MINT, TEAL),
        ("3", "ATIVO DE MENSAGERIA", "Número dedicado ou elegível • WABA • Phone Number ID • forma de pagamento", GREEN_LIGHT, GREEN_DARK),
        ("4", "APLICATIVO E WEBHOOK", "Meta App • token de acesso • App Secret • callback HTTPS • assinatura messages", AMBER_LIGHT, AMBER),
        ("5", "MODELOS E ESCALA", "Templates aprovados • teste controlado • qualidade • limites e custos monitorados no painel", RED_LIGHT, RED),
    ]
    y = 200
    for number, heading, body, fill, outline in rows:
        rounded_box(draw, (120, y, 1680, y + 130), fill, outline, 20, 3)
        draw.ellipse((150, y + 25, 230, y + 105), fill=hex_rgb(outline))
        text_center(draw, (150, y + 25, 230, y + 105), number, get_font(32, True), WHITE)
        draw.text((270, y + 25), heading, font=get_font(27, True), fill=hex_rgb(INK))
        draw.text((270, y + 70), body, font=get_font(22), fill=hex_rgb(GRAY))
        if number != "5":
            arrow(draw, (900, y + 130), (900, y + 160), TEAL_DARK, 5)
        y += 155
    draw.text(
        (120, 925),
        "Os valores, limites e requisitos podem mudar; o painel da conta e a documentação oficial são a fonte de verdade.",
        font=get_font(23, True),
        fill=hex_rgb(TEAL_DARK),
    )
    path = ASSETS / "figura-meta-onboarding.png"
    image.save(path, quality=95)
    out["meta"] = path

    return out


def remove_after_paragraph(document: Document, keep_count: int) -> None:
    body = document._element.body
    paragraphs_seen = 0
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        if tag == "p":
            if paragraphs_seen >= keep_count:
                body.remove(child)
            paragraphs_seen += 1
        elif paragraphs_seen >= keep_count:
            body.remove(child)


def set_run_font(run, size: float = 12, bold: bool | None = None, italic: bool | None = None, color: str = "000000"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def replace_paragraph_text(paragraph, text: str, size: float | None = None, bold: bool | None = None):
    while paragraph.runs:
        paragraph._element.remove(paragraph.runs[0]._element)
    run = paragraph.add_run(text)
    set_run_font(run, size or 12, bold=bold)
    return paragraph


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for style_name, size, bold, before, after in [
        ("Heading 1", 14, True, 22, 10),
        ("Heading 2", 12, True, 18, 6),
        ("Heading 3", 12, True, 14, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    styles["Heading 3"].font.italic = True

    if "NF Caption" not in styles:
        caption = styles.add_style("NF Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["NF Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing = 1.0

    if "NF Compact" not in styles:
        compact = styles.add_style("NF Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        compact = styles["NF Compact"]
    compact.font.name = "Times New Roman"
    compact._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    compact.font.size = Pt(10.5)
    compact.paragraph_format.space_after = Pt(2)
    compact.paragraph_format.line_spacing = 1.05
    compact.paragraph_format.first_line_indent = Cm(0)
    compact.paragraph_format.widow_control = True


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: Inches) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(round(width.inches * 1440))))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[float], indent: int = 110) -> None:
    total_twips = sum(int(round(width * 1440)) for width in widths)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(grid_col)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[float] | None = None,
    font_size: float = 9.5,
) -> object:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [5.8 / len(headers)] * len(headers)
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        set_cell_shading(cell, "DFF5F0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_width(cell, Inches(widths[idx]))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, font_size, bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            set_cell_width(cell, Inches(widths[idx]))
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, font_size)
    document.add_paragraph()
    return table


def add_caption(document: Document, text: str, source: str = "Elaboração própria (2026).") -> None:
    p = document.add_paragraph(style="NF Caption")
    p.add_run(text)
    for run in p.runs:
        set_run_font(run, 10, bold=True)
    s = document.add_paragraph(style="NF Caption")
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(8)
    s.add_run(f"Fonte: {source}")
    for run in s.runs:
        set_run_font(run, 9.5)


def add_picture(
    document: Document,
    path: Path,
    width: float,
    caption: str,
    source: str = "Elaboração própria (2026).",
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(document, caption, source)


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        r1 = p.add_run(first)
        set_run_font(r1, 12, bold=True)
        r2 = p.add_run(rest)
        set_run_font(r2, 12)
    else:
        r = p.add_run(text)
        set_run_font(r, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.widow_control = True


def add_compact(document: Document, text: str, bold_lead: str | None = None) -> None:
    p = document.add_paragraph(style="NF Compact")
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, 10.5, bold=True)
        r = p.add_run(text[len(bold_lead) :])
        set_run_font(r, 10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, 10.5)
    return p


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    p = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Cm(0.9 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.widow_control = True
    for run in p.runs:
        set_run_font(run, 11)
    if not p.runs:
        run = p.add_run(text)
        set_run_font(run, 11)
    else:
        p.text = text
        for run in p.runs:
            set_run_font(run, 11)


def add_numbered(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.widow_control = True
    p.text = text
    for run in p.runs:
        set_run_font(run, 11)


def add_heading(
    document: Document,
    text: str,
    level: int = 1,
    *,
    page_break_before: bool | None = None,
) -> None:
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    for run in p.runs:
        set_run_font(run, 14 if level == 1 else 12, bold=True, italic=(level == 3))
    if page_break_before is None:
        page_break_before = level == 1
    p.paragraph_format.page_break_before = page_break_before


def add_subheading(document: Document, text: str, level: int = 2) -> None:
    add_heading(document, text, level)


def add_callout(document: Document, title: str, body: str, fill: str = "E8FAF6", border: str = "1FB7A6") -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [5.85], indent=170)
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, Inches(5.85))
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 170, 150, 170)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), border)
        borders.append(el)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 11, bold=True, color=border)
    p2 = cell.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(body)
    set_run_font(r, 10.5)
    document.add_paragraph()


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Atualize este sumário no Word ou Google Docs, se necessário."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_page_break(document: Document) -> None:
    document.add_page_break()


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def patch_header_title(document: Document) -> None:
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                if "Colocar o Título" in paragraph.text or "TÍTULO DO TRABALHO" in paragraph.text:
                    replace_paragraph_text(paragraph, SHORT_TITLE, size=9)


def set_core_properties(document: Document) -> None:
    props = document.core_properties
    props.title = TITLE.title()
    props.subject = "Gate 4 — Testes e Validação do MVP Notify Flow"
    props.author = AUTHOR
    props.last_modified_by = "Codex, sob direção do autor"
    props.keywords = "Notify Flow; WhatsApp Cloud; Telegram; Gmail; BullMQ; Redis; LGPD; MVP"
    props.comments = (
        "Relatório técnico-científico gerado a partir do template institucional e "
        "das evidências verificáveis do repositório."
    )


def cover_and_title_page(document: Document) -> None:
    paragraphs = document.paragraphs
    replacements = {
        11: AUTHOR.upper(),
        12: "",
        13: "",
        19: TITLE,
        29: "Relatório de Impacto — Gate 4: Testes e Validação",
        35: "Goiânia",
        36: "2026",
        46: AUTHOR.upper(),
        47: "",
        48: "",
        54: TITLE,
        57: (
            "Relatório de Impacto apresentado ao Programa de Pós-Graduação Lato Sensu "
            "Especialização em Residência em Tecnologia da Informação do Instituto de "
            "Informática da Universidade Federal de Goiás, como requisito parcial para "
            "avaliação do Gate 4 — Testes e Validação do Projeto de Aplicação."
        ),
        59: f"Orientador: Prof. {ADVISOR}",
    }
    for index, text in replacements.items():
        replace_paragraph_text(paragraphs[index], text)
    for index in (11, 19, 29, 46, 54):
        paragraphs[index].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphs[index].paragraph_format.first_line_indent = Cm(0)
        for run in paragraphs[index].runs:
            set_run_font(run, 12 if index in (11, 29, 46) else 14, bold=True)
    for index in (35, 36):
        paragraphs[index].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphs[index].paragraph_format.first_line_indent = Cm(0)
    paragraphs[57].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraphs[57].paragraph_format.left_indent = Cm(7.5)
    paragraphs[57].paragraph_format.first_line_indent = Cm(0)
    paragraphs[57].paragraph_format.line_spacing = 1.0
    for run in paragraphs[57].runs:
        set_run_font(run, 10)
    paragraphs[59].alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraphs[59].paragraph_format.left_indent = Cm(7.5)
    paragraphs[59].paragraph_format.first_line_indent = Cm(0)
    for run in paragraphs[59].runs:
        set_run_font(run, 10)


def add_front_matter(document: Document) -> None:
    add_heading(document, "RESUMO EXECUTIVO", 1)
    add_body(
        document,
        "Este relatório apresenta o impacto, os testes e a validação do Notify Flow, um produto "
        "mínimo viável para administrar notificações consentidas por WhatsApp Cloud API, Telegram "
        "Bot API e e-mail. O problema enfrentado é a fragmentação entre canais, a dificuldade de "
        "comprovar consentimento e, sobretudo, o risco operacional de iniciar mensagens pelo "
        "WhatsApp sem observar as regras da API oficial e dos modelos aprovados pela Meta. A "
        "solução centraliza credenciais, contatos, convites, templates, conjuntos multicanal, "
        "campanhas, conversas e trilhas de auditoria, mantendo autorizações independentes por canal."
    )
    add_body(
        document,
        "A validação foi realizada por inspeção arquitetural, testes automatizados de unidade e "
        "contrato, análise estática, build de produção e verificação da configuração Docker Compose. "
        "Em 31 de julho de 2026, a API concluiu 305 testes e o frontend concluiu 169, totalizando "
        "474 testes aprovados, sem falhas. O lint da API, o build Vite do frontend e a validação "
        "estática do Compose também foram aprovados. Os fluxos reais dos provedores foram "
        "experimentados pelo autor durante a construção do MVP; entretanto, não houve estudo "
        "formal de usabilidade com amostra externa, o que é assumido como limitação metodológica."
    )
    add_body(
        document,
        "Os resultados indicam que o MVP está funcional para demonstração preliminar e que sua "
        "arquitetura reduz falhas em cascata: cada entrega possui estado próprio, tentativas com "
        "backoff exponencial, logs por contato e canal e reconciliação de recibos. O uso do "
        "WhatsApp Cloud substitui a dependência de automações não oficiais para notificações "
        "iniciadas pela organização, reduzindo o risco de bloqueio associado ao uso inadequado do "
        "WhatsApp Web. O impacto principal é dar ao contato o controle da autorização e, ao "
        "administrador, visibilidade operacional para enviar apenas por destinos identificados e "
        "permitidos."
    )
    add_compact(
        document,
        "Palavras-chave: notificações multicanal; WhatsApp Cloud API; Telegram Bot API; consentimento; "
        "LGPD; BullMQ; Redis; webhooks; observabilidade.",
    )

    add_heading(document, "LISTA DE FIGURAS", 1)
    for item in [
        "Figura 1 — Arquitetura ponta a ponta do Notify Flow",
        "Figura 2 — Jornada de consentimento e vínculo de identidades",
        "Figura 3 — Área pública Meu Perfil em desktop e dispositivo móvel",
        "Figura 4 — Preparação do WhatsApp Cloud para produção",
        "Figura 5 — Ciclo de entrega confiável, tentativas e auditoria",
    ]:
        add_compact(document, item)

    add_heading(document, "LISTA DE QUADROS E TABELAS", 1)
    for item in [
        "Quadro 1 — Perfis e responsabilidades",
        "Quadro 2 — Componentes da arquitetura",
        "Quadro 3 — Regras de autorização por canal",
        "Quadro 4 — Requisitos para produção no WhatsApp Cloud",
        "Quadro 5 — Variáveis de ambiente por domínio",
        "Tabela 1 — Evidências de testes e validações",
        "Tabela 2 — Riscos, controles e limitações",
        "Tabela 3 — Histórias de usuário do administrador",
        "Tabela 4 — Histórias de usuário do contato",
    ]:
        add_compact(document, item)

    add_heading(document, "LISTA DE SIGLAS E ABREVIATURAS", 1)
    add_table(
        document,
        ["Sigla", "Significado"],
        [
            ("API", "Interface de Programação de Aplicações"),
            ("CNPJ", "Cadastro Nacional da Pessoa Jurídica"),
            ("DTO", "Objeto de Transferência de Dados"),
            ("HMAC", "Código de Autenticação de Mensagem com Hash"),
            ("JWT", "JSON Web Token"),
            ("LGPD", "Lei Geral de Proteção de Dados Pessoais"),
            ("MVP", "Produto Mínimo Viável"),
            ("SMTP", "Protocolo Simples de Transferência de Correio"),
            ("TTL", "Tempo de vida de um registro ou chave"),
            ("WABA", "Conta do WhatsApp Business"),
            ("WSS", "WebSocket Seguro"),
        ],
        widths=[1.0, 4.8],
    )

    add_heading(document, "SUMÁRIO", 1)
    add_toc(document)
    add_callout(
        document,
        "Nota de compatibilidade",
        "O arquivo foi preparado em DOCX para edição no Microsoft Word e importação no Google Docs. "
        "Ao abrir o documento, atualize o sumário automático caso o editor não o faça sozinho.",
        fill="F5FAF9",
        border="087E73",
    )


def add_introduction(document: Document) -> None:
    add_heading(document, "1 INTRODUÇÃO", 1, page_break_before=False)
    add_subheading(document, "1.1 Contexto e problema", 2)
    add_body(
        document,
        "Organizações que precisam avisar clientes, servidores ou participantes de programas "
        "frequentemente combinam mensageria instantânea e e-mail. Na prática, cada provedor possui "
        "identificadores, políticas, limites, modelos e formas de comprovar entrega diferentes. "
        "Quando esses detalhes são tratados por planilhas, scripts isolados ou sessões não oficiais, "
        "o administrador perde a visão de quem autorizou cada canal, por qual campanha a pessoa foi "
        "incluída e por que uma entrega falhou."
    )
    add_body(
        document,
        "O WhatsApp concentra a dor mais sensível. Fora da janela de atendimento iniciada pelo "
        "cliente, a organização deve utilizar modelos aprovados na conta remetente. A automação de "
        "uma sessão comum do WhatsApp Web não equivale à API oficial e, quando empregada para "
        "disparos, amplia o risco de restrição do número e não oferece as mesmas garantias de "
        "governança. O Notify Flow foi concebido para trabalhar com o WhatsApp Cloud API e deixar "
        "a conversa livre restrita à janela permitida, mantendo o Telegram e o e-mail como canais "
        "complementares."
    )
    add_body(
        document,
        "O segundo problema é o consentimento. A simples existência de um telefone, endereço de "
        "e-mail ou nome de usuário não prova que a pessoa deseja receber notificações. O sistema "
        "precisa capturar a interação real, vincular identidades sem duplicar contatos, registrar "
        "a origem da autorização, possibilitar revogação e impedir que uma entrega não autorizada "
        "seja silenciosamente tratada como sucesso."
    )

    add_subheading(document, "1.2 Objetivo geral e objetivos específicos", 2)
    add_body(
        document,
        "O objetivo geral é demonstrar a viabilidade de uma plataforma web responsiva que centraliza "
        "notificações multicanal consentidas, com filas resilientes, logs auditáveis, integrações "
        "oficiais e recursos de autogestão para o contato."
    )
    for item in [
        "separar as configurações e permissões de WhatsApp Cloud, Telegram e e-mail;",
        "permitir a criação amigável de templates por canal e de conjuntos multicanais;",
        "registrar contatos por interação, convite ou cadastro administrativo sem criar destinos artificiais;",
        "processar campanhas em fila, isolando a falha de um destinatário dos demais;",
        "exibir mensagens, recibos, tentativas e erros em tempo real sem expor segredos;",
        "oferecer ao contato uma área segura para consultar dados, vínculos e histórico;",
        "documentar a execução local em Docker e a implantação por Blueprint no Render;",
        "validar o MVP por testes automatizados e critérios de aceitação reproduzíveis.",
    ]:
        add_bullet(document, item)

    add_subheading(document, "1.3 Escopo do MVP", 2)
    add_body(
        document,
        "O escopo abrange painel administrativo, página pública de convite, área Meu Perfil, API "
        "REST, comunicação em tempo real, persistência MongoDB, fila BullMQ sobre Redis, webhooks "
        "do WhatsApp Cloud e Telegram, SMTP do Gmail, Docker Compose e Blueprint do Render. O MVP "
        "não pretende substituir uma central completa de atendimento, um provedor de CRM ou a "
        "governança nativa da Meta. Também não aprova modelos no WhatsApp Manager: o administrador "
        "cadastra no Notify Flow o nome e a estrutura de um modelo que já existe na conta remetente."
    )
    add_body(
        document,
        "A solução atual não usa WhatsApp Web JS para iniciar notificações. As conversas exibidas no "
        "painel WhatsApp são constituídas pelos eventos recebidos no webhook oficial e pelas respostas "
        "enviadas pela Graph API. A plataforma também não importa um histórico completo da Meta; ela "
        "mantém o histórico que efetivamente atravessa sua própria integração."
    )

    add_subheading(document, "1.4 Relevância e impacto esperado", 2)
    add_body(
        document,
        "O valor do projeto não está apenas em enviar mensagens. Está em tornar explícito o caminho "
        "entre consentimento, identidade, template, fila, provedor e recibo. Esse encadeamento permite "
        "que o administrador explique por que determinado contato recebeu — ou deixou de receber — "
        "uma comunicação, ao mesmo tempo em que oferece ao titular mecanismos de consulta e revogação."
    )
    add_callout(
        document,
        "Hipótese de impacto",
        "Ao concentrar notificações na API oficial do WhatsApp e registrar consentimento por canal, "
        "o Notify Flow reduz risco de bloqueio por automação inadequada, retrabalho operacional e "
        "envios indevidos; aumenta, ainda, a rastreabilidade de campanhas e a autonomia do contato.",
    )


def add_theory(document: Document) -> None:
    add_heading(document, "2 FUNDAMENTAÇÃO TEÓRICA", 1)
    add_subheading(document, "2.1 Mensageria oficial, janela de atendimento e modelos", 2)
    add_body(
        document,
        "A WhatsApp Business Platform diferencia mensagens de serviço, enviadas dentro de uma janela "
        "de atendimento iniciada pelo usuário, e mensagens iniciadas pela empresa. A janela dura 24 "
        "horas e é renovada quando o cliente envia uma nova mensagem. Nesse período, respostas livres "
        "de serviço não são cobradas segundo a página de preços consultada; fora dele, a empresa deve "
        "usar um modelo aprovado e está sujeita à cobrança por mensagem entregue, de acordo com "
        "mercado e categoria. As categorias atuais incluem marketing, utilidade, autenticação e "
        "serviço (META, 2026a)."
    )
    add_body(
        document,
        "O modelo aprovado é um contrato entre a conta remetente e a Meta. Seu nome, idioma, "
        "componentes e parâmetros precisam corresponder ao artefato cadastrado no WhatsApp Manager. "
        "Por isso, o Notify Flow oferece formulários que montam o payload sem exigir JSON do usuário, "
        "mas não presume aprovação. Os itens de exemplo vinculados ao número de teste ou ao número de "
        "produção só funcionam quando o modelo homônimo está disponível para aquele remetente."
    )
    add_body(
        document,
        "No Telegram, o bot pode responder a pessoas que iniciaram ou autorizaram a conversa. Ele não "
        "descobre o telefone arbitrariamente: o número só é confiável quando o próprio usuário "
        "compartilha o contato e o identificador do contato corresponde ao remetente. Já o e-mail "
        "utiliza SMTP autenticado e exige endereço válido, permissão ativa e, nos fluxos de atualização "
        "pelo chat, confirmação por código temporário."
    )

    add_subheading(document, "2.2 Consentimento, minimização e direitos do titular", 2)
    add_body(
        document,
        "A LGPD exige finalidade, adequação, necessidade, transparência, segurança e mecanismos para "
        "o exercício de direitos. No contexto do MVP, isso se traduz em autorizações independentes, "
        "termos apresentados na jornada pública, registro da fonte da decisão e proibição de inventar "
        "uma identidade de destino apenas porque outro canal foi autorizado. O contato pode consultar "
        "seus dados e entregas, revogar um canal, remover vínculos próprios com convites e sair de "
        "grupos mediante confirmação."
    )
    add_body(
        document,
        "Os dados sensíveis de configuração são criptografados com AES-256-GCM e os identificadores "
        "pesquisáveis usam índices cegos baseados em HMAC. A exclusão do contato pseudonimiza registros "
        "históricos necessários à auditoria, em vez de prometer apagamento indiscriminado. Essa "
        "distinção é importante: a plataforma reduz exposição e preserva comprovação operacional, "
        "mas não implementa criptografia de ponta a ponta entre o administrador e os provedores."
    )

    add_subheading(document, "2.3 Processamento assíncrono e observabilidade", 2)
    add_body(
        document,
        "Uma campanha multicanal é naturalmente sujeita a falhas parciais. O endereço pode estar "
        "inválido, o modelo pode não existir na conta, a janela do WhatsApp pode estar fechada, um "
        "token pode expirar ou o provedor pode responder com limitação temporária. Processar todos os "
        "destinatários na mesma requisição HTTP prolongaria o tempo de resposta e permitiria que uma "
        "falha interrompesse o conjunto."
    )
    add_body(
        document,
        "BullMQ materializa cada entrega como um trabalho em Redis. O worker pode operar com "
        "concorrência controlada, repetir falhas transitórias e preservar o estado do restante da "
        "campanha. No Notify Flow, o worker usa concorrência cinco, até quatro tentativas e backoff "
        "exponencial a partir de dois segundos. O MongoDB mantém o registro durável da entrega e um "
        "processo de recuperação procura itens pendentes de enfileiramento. A semântica externa deve "
        "ser descrita como pelo menos uma vez em cenários ambíguos: idempotência e travas reduzem "
        "duplicidade local, mas nenhum sistema distribuído deve prometer exatamente uma entrega "
        "quando o provedor aceitou a mensagem e a resposta se perdeu."
    )
    add_body(
        document,
        "Observabilidade significa correlacionar requisição, campanha, entrega e recibo sem registrar "
        "segredos. A API atribui requestId, redige campos sensíveis e mantém resultados por contato e "
        "canal. WebSockets transportam atualizações incrementais para o painel; webhooks transportam "
        "eventos dos provedores para a API. São mecanismos complementares: webhook é servidor a "
        "servidor; WebSocket mantém a interface conectada ao backend."
    )

    add_subheading(document, "2.4 Contêineres e infraestrutura como código", 2)
    add_body(
        document,
        "O Docker Compose descreve um ambiente local composto por MongoDB, Redis, API e frontend "
        "Nginx. O Blueprint do Render cumpre papel semelhante em produção, mas não replica o Compose "
        "literalmente: o frontend é um serviço web público, a API é um serviço privado e o Redis é "
        "um Key Value gerenciado. O MongoDB é fornecido externamente pelo Atlas. Variáveis não "
        "sensíveis podem ser declaradas no YAML; segredos devem ser gerados ou solicitados durante "
        "a sincronização, nunca versionados."
    )


def add_methodology(document: Document) -> None:
    add_heading(document, "3 METODOLOGIA", 1)
    add_subheading(document, "3.1 Estratégia de desenvolvimento", 2)
    add_body(
        document,
        "O trabalho foi desenvolvido individualmente por Samuel Victor Oliveira Lima, com orientação "
        "de Leonardo Oliveira, por ciclos incrementais. Cada ciclo partiu de um fluxo observável — "
        "por exemplo, receber um webhook, autorizar um canal, criar uma campanha ou abrir o Meu Perfil — "
        "e foi refinado até que regras de negócio, interface, persistência e testes convergissem. O "
        "repositório separa `api` e `frontend`, possui documentação própria em cada diretório e mantém "
        "a infraestrutura local e de produção na raiz."
    )
    add_body(
        document,
        "A análise para este Gate utilizou quatro fontes de evidência: inspeção do código e dos "
        "artefatos de infraestrutura; execução das suítes automatizadas; build e análise estática; e "
        "testes manuais conduzidos pelo autor nos painéis de Meta, Telegram e Gmail durante a "
        "implementação. Nenhum dado de credencial foi reproduzido neste documento."
    )

    add_subheading(document, "3.2 Perfis e histórias de usuário", 2)
    add_table(
        document,
        ["Perfil", "Responsabilidade", "Resultado esperado"],
        [
            (
                "Administrador",
                "Configurar canais, contatos, templates, convites, campanhas e documentos.",
                "Operar envios com rastreabilidade, sem que uma falha paralise os demais.",
            ),
            (
                "Contato",
                "Iniciar a interação, consentir, consultar dados, vínculos e entregas.",
                "Controlar seus canais e receber apenas comunicações autorizadas.",
            ),
            (
                "Provedor",
                "Validar credenciais, aceitar mensagens e devolver eventos/recibos.",
                "Aplicar políticas próprias; o Notify Flow não as contorna.",
            ),
        ],
        widths=[1.1, 2.35, 2.35],
    )
    add_caption(document, "Quadro 1 — Perfis e responsabilidades")
    add_body(
        document,
        "As histórias completas estão no Apêndice A. Elas foram convertidas em critérios verificáveis, "
        "como impedir resposta livre fora da janela de 24 horas, ignorar um destinatário sem consentimento "
        "sem abortar a campanha e exigir confirmação para revogar vínculo ou permissão."
    )

    add_subheading(document, "3.3 Estratégia de testes e critérios de aceitação", 2)
    add_body(
        document,
        "A API foi testada com o test runner nativo do Node, isolando managers, controllers, "
        "middlewares, validações, criptografia, filas, webhooks e políticas de consentimento. O frontend "
        "foi testado com Vitest, cobrindo serviços HTTP, autenticação, sockets, contatos, convites, "
        "perfil, templates, campanhas, diálogos e navegação. O lint verificou consistência da API; o "
        "build Vite demonstrou que os módulos do frontend são empacotáveis para produção; e `docker "
        "compose config --quiet` validou a composição estática."
    )
    for item in [
        "todas as suítes automatizadas devem encerrar sem falha;",
        "nenhum segredo real deve aparecer no código, nos logs ou no relatório;",
        "o webhook do WhatsApp deve validar challenge e assinatura HMAC;",
        "o webhook do Telegram deve validar segredo e deduplicar update_id;",
        "a campanha deve continuar quando um destinatário falha ou não autorizou o canal;",
        "a resposta livre do WhatsApp deve respeitar a janela de 24 horas;",
        "o contato deve poder revogar e consultar as próprias permissões;",
        "a versão de produção do frontend deve ser compilada com sucesso.",
    ]:
        add_bullet(document, item)

    add_subheading(document, "3.4 Limitações metodológicas", 2)
    add_body(
        document,
        "Os testes automatizados são predominantemente unitários e de contrato, com dependências "
        "externas simuladas. Eles não substituem um teste ponta a ponta contínuo contra Atlas, Redis, "
        "SMTP, Telegram, Graph API e Render. Também não foi realizada pesquisa de usabilidade com "
        "participantes externos nem comparação estatística antes/depois. Assim, o impacto é analisado "
        "por evidência funcional, redução de risco arquitetural e capacidade de auditoria; não como "
        "ganho quantitativo já comprovado em uma organização."
    )


def add_development_results(document: Document, diagrams: dict[str, Path]) -> None:
    add_heading(document, "4 DESENVOLVIMENTO, RESULTADOS E DISCUSSÃO", 1)
    add_subheading(document, "4.1 Visão geral da solução", 2)
    add_body(
        document,
        "O Notify Flow é uma aplicação web responsiva composta por três superfícies. A primeira é o "
        "painel administrativo, no qual se configuram canais, contatos, templates, conjuntos, convites, "
        "campanhas e documentos legais. A segunda é a camada pública de convite, que apresenta termos "
        "e direciona o visitante ao canal selecionado. A terceira é o Meu Perfil, área em que o contato "
        "consulta seus dados, permissões, convites, grupos e histórico de entregas."
    )
    add_picture(
        document,
        diagrams["architecture"],
        5.9,
        "Figura 1 — Arquitetura ponta a ponta do Notify Flow",
    )
    add_table(
        document,
        ["Componente", "Tecnologia", "Responsabilidade principal"],
        [
            ("Frontend", "Vue 3, Quasar, Pinia, Vue Router, Axios", "Interface responsiva, sessões separadas, preview e administração."),
            ("Tempo real", "Socket.IO", "Eventos, mensagens, recibos, logs e central administrativa."),
            ("API", "Node.js, Express, Zod", "Rotas, validação, autorização, políticas e orquestração."),
            ("Dados", "MongoDB, Mongoose", "Contatos, consentimentos, campanhas, eventos e configurações criptografadas."),
            ("Fila", "Redis/Valkey, BullMQ", "Concorrência, tentativas, backoff e recuperação de entregas."),
            ("Provedores", "Meta Graph API, Telegram Bot API, Gmail SMTP", "Entrega e recepção de eventos externos."),
            ("Infraestrutura", "Docker Compose, Nginx, Render Blueprint", "Execução local e implantação reproduzível."),
        ],
        widths=[1.05, 1.85, 2.9],
        font_size=9.0,
    )
    add_caption(document, "Quadro 2 — Componentes da arquitetura")

    add_subheading(document, "4.2 Fluxo orientado ao consentimento", 2)
    add_body(
        document,
        "O contato pode chegar por convite, comando ou cadastro administrativo. Um clique de convite "
        "gera rastreabilidade, mas não autoriza sozinho o destino: a identidade real precisa ser "
        "observada pelo webhook ou confirmada por um mecanismo seguro. O WhatsApp e o Telegram "
        "reconhecem comandos configuráveis; o e-mail digitado em qualquer momento do chat inicia um "
        "desafio de validação com seis dígitos, válido por quinze minutos e reenviável após dois minutos."
    )
    add_picture(
        document,
        diagrams["consent"],
        5.9,
        "Figura 2 — Jornada de consentimento e vínculo de identidades",
    )
    add_table(
        document,
        ["Canal", "Como a identidade é confirmada", "Como autoriza", "Como revoga"],
        [
            ("WhatsApp Cloud", "wa_id/telefone observado no webhook oficial.", "Comando dinâmico, Meu Perfil ou decisão administrativa confirmada.", "Meu Perfil ou painel administrativo com confirmação."),
            ("Telegram", "chat_id recebido; telefone apenas por contact.user_id igual ao remetente.", "Comando/invite assinado e interação com o bot.", "Comando de parada, Meu Perfil ou administração."),
            ("E-mail", "Código enviado ao endereço informado e validado no chat/perfil.", "Ativado após validação ou confirmação no Meu Perfil.", "Meu Perfil ou administração."),
        ],
        widths=[1.05, 2.05, 1.7, 1.1],
        font_size=8.8,
    )
    add_caption(document, "Quadro 3 — Regras de autorização por canal")
    add_body(
        document,
        "O vínculo entre Telegram e WhatsApp só é consolidado quando existe evidência comum, como o "
        "telefone compartilhado pelo próprio usuário. O sistema consulta índices cegos para localizar "
        "um contato existente, transfere a identidade e evita criar perfis paralelos. Se outro canal "
        "ainda não possui identidade real, a autorização pode permanecer pendente, mas o sistema não "
        "fabrica um chat_id ou wa_id."
    )

    add_subheading(document, "4.3 Funcionalidades do administrador", 2)
    add_subheading(document, "4.3.1 Início e configurações", 3)
    add_body(
        document,
        "A tela inicial salva cada provedor separadamente. O WhatsApp recebe token, Phone Number ID, "
        "Business Account ID, número público, versão da Graph API, App Secret e verify token; o "
        "Telegram recebe token e URL pública, identifica o bot por getMe e registra o webhook; o Gmail "
        "recebe conta, senha de aplicativo e remetente. Os campos persistidos aparecem mascarados e só "
        "podem ser revelados por administrador autenticado, em rota sem cache, limitada e auditada. "
        "Também são configuráveis os comandos, respostas amigáveis, links úteis e textos de autorização."
    )
    add_subheading(document, "4.3.2 Contatos, convites e grupos", 3)
    add_body(
        document,
        "O administrador busca, cadastra, edita e pseudonimiza contatos, acompanha identidades técnicas "
        "e permissões e cria grupos. Convites possuem slug, identidade visual, QR Code e links por canal. "
        "Quando o contato chega por um convite, o vínculo é armazenado; a sincronização por convite "
        "cria ou atualiza grupos sem duplicar participantes. Um mesmo contato pode ter mais de um convite "
        "e, por isso, participar de mais de um segmento."
    )
    add_subheading(document, "4.3.3 Templates e conjuntos multicanais", 3)
    add_body(
        document,
        "A biblioteca oferece editores específicos. WhatsApp usa nome oficial, idioma, componentes e "
        "parâmetros; Telegram suporta texto, mídia validada e menus; e-mail suporta texto ou HTML "
        "sanitizado. Um conjunto agrega de um a três templates, no máximo um por canal, e pode ser "
        "associado a um convite. O mesmo template pode participar de vários conjuntos. Essa modelagem "
        "reduz divergência entre campanhas recorrentes e evita que o administrador manipule payloads "
        "JSON manualmente."
    )
    add_subheading(document, "4.3.4 Notificações, canais e governança", 3)
    add_body(
        document,
        "A campanha pode usar um conjunto, templates isolados ou uma mensagem rápida quando a política "
        "do canal permite. O administrador escolhe um, dois ou três canais, contatos ou grupos, preenche "
        "variáveis e revisa uma prévia antes de confirmar. As páginas de WhatsApp, Telegram e Gmail "
        "reúnem disparos específicos; WhatsApp e Telegram também mostram conversas em tempo real. "
        "Governança inclui termos e privacidade versionados, auditoria de links temporários e central de "
        "eventos administrativos com retenção de trinta dias."
    )

    add_subheading(document, "4.4 Funcionalidades do contato e Meu Perfil", 2)
    add_body(
        document,
        "O acesso do contato é separado da sessão administrativa. O link temporário é assinado, de uso "
        "único e possui validade configurável de até sete dias; desafios de código possuem tentativas e "
        "limites de reenvio. Os comandos `/login` e `/meu-perfil` podem gerar o link no WhatsApp ou "
        "Telegram. No perfil, o contato edita dados, ativa ou revoga e-mail, consulta identidades dos "
        "canais, remove seus próprios vínculos com convites, sai de grupos e acompanha apenas as "
        "entregas destinadas a ele."
    )
    if PROFILE_IMAGE.exists():
        add_picture(
            document,
            PROFILE_IMAGE,
            5.9,
            "Figura 3 — Área pública Meu Perfil em desktop e dispositivo móvel",
            "Arte visual do projeto Notify Flow (2026).",
        )
    add_body(
        document,
        "A separação de sessões limita a superfície de privilégio: o link do contato não concede acesso "
        "ao painel administrativo. Da mesma forma, desativar um canal não elimina automaticamente o "
        "histórico já pseudonimizado, mas impede novas entregas naquele destino. A interface explicita "
        "convites e grupos para que o titular compreenda por que pode ser selecionado em uma campanha."
    )

    add_subheading(document, "4.5 Integração com WhatsApp Cloud e Meta", 2)
    add_subheading(document, "4.5.1 Preparação empresarial e burocrática", 3)
    add_body(
        document,
        "A implantação produtiva começa fora do código. A organização precisa de empresa real com "
        "dados consistentes, conta do Facebook protegida por autenticação em dois fatores, Página e "
        "portfólio empresarial no Business Manager/Meta Business Suite. A Meta pode solicitar atos "
        "constitutivos, licença, documento fiscal ou conta de serviço, todos legíveis, válidos e "
        "compatíveis com o nome, endereço e telefone informados. Não existe um prazo universal de "
        "aprovação; portanto, a verificação deve ser tratada como dependência de cronograma, e não como "
        "atividade concluída no dia da entrega."
    )
    add_body(
        document,
        "Para reduzir risco operacional, recomenda-se um número dedicado, novo ou confirmado como "
        "elegível para o fluxo disponível na conta. A afirmação de que todo número precisa ser "
        "absolutamente inédito seria excessiva, pois migração e coexistência dependem dos recursos "
        "oferecidos pela Meta; entretanto, reutilizar um número já preso a outra conta sem preparar a "
        "migração é uma causa comum de falha de registro. O painel da conta é a autoridade sobre "
        "elegibilidade."
    )
    add_picture(
        document,
        diagrams["meta"],
        5.9,
        "Figura 4 — Preparação do WhatsApp Cloud para produção",
    )
    add_table(
        document,
        ["Etapa", "Evidência necessária", "Risco se omitida"],
        [
            ("Empresa e portfólio", "CNPJ e dados públicos consistentes.", "Bloqueio da verificação ou limitação do portfólio."),
            ("Segurança", "2FA, administradores legítimos e acesso mínimo.", "Tomada de conta ou rejeição de ativos."),
            ("Número", "Recebimento de SMS/ligação e elegibilidade.", "Impossibilidade de registrar o remetente."),
            ("Aplicativo", "Meta App, produto WhatsApp e permissões.", "Token ou webhook sem acesso ao WABA."),
            ("Webhook", "URL HTTPS, verify token e App Secret.", "Eventos rejeitados ou sem autenticidade."),
            ("Pagamento", "Forma válida para mensagens cobradas.", "Bloqueio de mensagens iniciadas pela empresa."),
            ("Modelos", "Nome/idioma aprovados na conta remetente.", "Erro de template inexistente ou não aprovado."),
        ],
        widths=[1.05, 2.45, 2.3],
        font_size=8.8,
    )
    add_caption(document, "Quadro 4 — Requisitos para produção no WhatsApp Cloud")

    add_subheading(document, "4.5.2 Credenciais, webhook e comandos", 3)
    add_body(
        document,
        "O Phone Number ID é o identificador técnico usado na rota `/messages`; o Business Account ID "
        "identifica a WABA; o access token autoriza a Graph API; o App Secret valida a assinatura "
        "`X-Hub-Signature-256`; e o verify token é um segredo escolhido pelo administrador para o "
        "challenge inicial do webhook. O número público é mantido separadamente porque serve a links "
        "`wa.me` e à apresentação ao usuário. Em produção, tokens devem ser permanentes ou renovados "
        "por processo seguro, nunca colados em repositórios ou documentos."
    )
    add_body(
        document,
        "Ao receber `/notify-me`, o fluxo registra o consentimento WhatsApp para identidades reais "
        "disponíveis e devolve instruções amigáveis. `/login` gera um link temporário de acesso ao Meu "
        "Perfil. Um e-mail válido digitado durante a conversa inicia sua validação; somente após o "
        "código correto o endereço é criado ou alterado e a permissão de e-mail é ativada. A resposta "
        "livre continua limitada à janela de 24 horas; a notificação iniciada pela organização usa "
        "modelo aprovado."
    )
    add_body(
        document,
        "A instalação mantém três exemplos protegidos contra exclusão: "
        "`jaspers_market_plain_text_v1` e `jaspers_market_order_confirmation_v1`, associados ao "
        "ambiente de número de teste da Meta, e `3p_direct_integration_test_template`, usado para "
        "validar a integração do número de produção em modo de teste. Eles não são modelos universais: "
        "o mesmo nome e idioma precisam existir e estar aprovados na WABA que atende o remetente. O "
        "antigo `verify_code_1` foi retirado da lista fixa; o login atual usa o comando `/login` e um "
        "link temporário assinado, sem depender desse modelo de autenticação."
    )

    add_subheading(document, "4.5.3 Custos, qualidade e limites", 3)
    add_body(
        document,
        "Na política consultada em julho de 2026, a cobrança da WhatsApp Business Platform ocorre por "
        "mensagem entregue e varia por mercado e categoria. Mensagens de serviço não são cobradas e "
        "mensagens de utilidade enviadas em resposta ao usuário dentro da janela de atendimento também "
        "podem ser gratuitas. Marketing, autenticação e utilidade fora da janela seguem as tabelas "
        "vigentes. Como preços e regras mudam, o relatório não fixa um valor em reais: a estimativa deve "
        "usar a página oficial e o país do destinatário no momento da contratação (META, 2026a)."
    )
    add_body(
        document,
        "A escala de mensagens iniciadas pela empresa depende do nível exibido no WhatsApp Manager, da "
        "verificação e da qualidade. O material oficial de onboarding descreve evolução por degraus "
        "como 1 mil, 10 mil, 100 mil destinatários únicos por período de 24 horas e, depois, ilimitado, "
        "mas contas de teste ou novas podem apresentar limites iniciais distintos. O painel da própria "
        "conta, e não um número reproduzido em documentação acadêmica, é a fonte de verdade. Bloqueios, "
        "denúncias e baixa qualidade podem impedir a expansão."
    )
    add_callout(
        document,
        "Discussão sobre a principal dor",
        "A API oficial não elimina toda possibilidade de restrição, porque políticas e qualidade ainda "
        "se aplicam. Ela, porém, fornece o caminho suportado para modelos, limites, recibos e cobrança. "
        "Isso é qualitativamente mais seguro e auditável do que automatizar uma sessão WhatsApp Web para "
        "disparos iniciados pela organização.",
        fill="E9FAF0",
        border="128C7E",
    )

    add_subheading(document, "4.6 Integração com Telegram", 2)
    add_body(
        document,
        "O bot é criado no `@BotFather`, que entrega o token. Esse token deve ser tratado como senha: "
        "qualquer pessoa que o possua controla o bot. Depois de salvo, o Notify Flow chama `getMe` para "
        "identificar nome e username. Para receber atualizações, registra uma URL HTTPS em "
        "`/api/webhooks/telegram` e envia um secret token; a API compara o cabeçalho correspondente de "
        "forma segura e deduplica `update_id` no Redis."
    )
    add_body(
        document,
        "A pessoa precisa iniciar o bot. `/notify-me` registra autorização; `/verify-me` abre o menu de "
        "vínculo, acesso ao Meu Perfil e ajuda; `/login` entrega um link temporário. Quando o usuário "
        "compartilha o próprio contato por botão nativo, o backend confirma que `contact.user_id` "
        "corresponde ao remetente antes de associar o telefone. Um e-mail digitado em qualquer ponto "
        "inicia o mesmo desafio de confirmação usado no WhatsApp."
    )
    add_body(
        document,
        "Templates Telegram podem enviar texto, foto, vídeo e menus. Links de mídia passam por validação "
        "HTTPS, bloqueio de endereços privados para reduzir SSRF, limite de redirecionamentos, timeout, "
        "verificação de assinatura do arquivo e limites de tamanho antes de chamar `sendPhoto` ou "
        "`sendVideo`. O bot não obtém silenciosamente o número de qualquer usuário e estar no mesmo "
        "grupo não autoriza mensagem privada."
    )

    add_subheading(document, "4.7 Integração com Gmail", 2)
    add_body(
        document,
        "O MVP usa Nodemailer com o serviço Gmail por SMTP. A conta Gmail autentica a conexão; a senha "
        "de aplicativo, normalmente com 16 caracteres, substitui a senha principal; o campo remetente "
        "define o endereço apresentado no `From`, que deve ser a própria conta ou um alias autorizado. "
        "SMTP cuida do transporte entre a API e os servidores do Google, tipicamente por TLS."
    )
    add_body(
        document,
        "Para gerar uma senha de aplicativo, a conta precisa de verificação em duas etapas. O recurso "
        "pode não aparecer em contas corporativas administradas, Advanced Protection ou configurações "
        "restritas a chaves de segurança. O Google revoga senhas de aplicativo quando a senha principal "
        "é alterada. Em uma evolução de produção com múltiplas organizações, OAuth 2.0 e a Gmail API "
        "seriam preferíveis por oferecerem escopo e revogação mais granulares; o MVP atual não "
        "implementa OAuth."
    )

    add_subheading(document, "4.8 Fila, tentativas, recibos e logs", 2)
    add_body(
        document,
        "Ao confirmar uma campanha, a API cria notificações e entregas no MongoDB antes de enviá-las ao "
        "BullMQ. Essa ordem evita que uma interrupção entre a requisição e o Redis apague a intenção de "
        "envio. O worker reivindica a entrega, renova um heartbeat, monta o payload próprio do canal e "
        "classifica falhas transitórias e permanentes. Itens transitórios recebem até quatro tentativas "
        "com backoff; itens sem permissão são ignorados com motivo explícito e não consomem envio."
    )
    add_picture(
        document,
        diagrams["queue"],
        5.9,
        "Figura 5 — Ciclo de entrega confiável, tentativas e auditoria",
    )
    add_body(
        document,
        "O WhatsApp fornece recibos como enviado, entregue, lido ou falhou; Telegram e SMTP devolvem "
        "aceitação técnica, mas não possuem necessariamente o mesmo nível de confirmação de leitura. "
        "Por isso, os status não devem ser interpretados como equivalentes entre canais. A interface "
        "mostra o resultado de cada contato e canal, número de tentativas e mensagem operacional "
        "sanitizada. Os logs gerais expiram por padrão em 180 dias; recibos em sete dias; eventos da "
        "central e conversas Cloud em 30 dias; backups de conversas em 90 dias."
    )
    add_body(
        document,
        "O Redis também deduplica webhooks e mantém estado efêmero. Em produção ele é obrigatório: a API "
        "falha fechada se não conseguir conectar, porque aceitar campanhas sem fila criaria falsa "
        "confiança. A política `noeviction` evita apagar trabalhos sob pressão de memória, mas exige "
        "monitoramento de capacidade."
    )

    add_subheading(document, "4.9 Execução local com Docker", 2)
    add_body(
        document,
        "Na raiz do projeto, o Docker Compose cria quatro serviços: MongoDB 7, Redis 7.4, API e frontend "
        "Nginx. Apenas o frontend e a API são publicados no loopback do host; Mongo e Redis permanecem "
        "na rede interna. Os volumes `mongo_data` e `redis_data` preservam dados entre recriações. A "
        "API possui health check que consulta Mongo e Redis; o frontend usa `/healthz` para verificar "
        "o Nginx."
    )
    add_numbered(document, "Instalar Docker Desktop e confirmar que o daemon está ativo.")
    add_numbered(document, "Copiar os arquivos `.env.example` adequados e preencher apenas segredos locais.")
    add_numbered(document, "Executar `docker compose config --quiet` para validar a composição.")
    add_numbered(document, "Executar `docker compose up --build -d` na raiz do projeto.")
    add_numbered(document, "Abrir `http://localhost:8080` e verificar `http://localhost:8080/api/health`.")
    add_numbered(document, "Acompanhar `docker compose logs -f api frontend` durante os testes.")
    add_body(
        document,
        "Quando `MONGODB_URI` não é fornecida no Compose, a intenção é usar o serviço local. Um placeholder "
        "de Atlas copiado literalmente deve ser removido ou substituído, caso contrário a API tentará "
        "resolver um endereço inexistente. Credenciais de provedores podem ser cadastradas pela interface "
        "e ficam criptografadas no banco."
    )

    add_subheading(document, "4.10 Implantação no Render por Blueprint", 2)
    add_body(
        document,
        "O arquivo `render.yaml` declara três recursos no Oregon, todos no plano `starter`: o web service "
        "`notify-flow`, o private service `api` e o Key Value `notify-flow-redis`. O frontend compila Vue "
        "e serve arquivos por Nginx. Como a API é privada, ela não possui URL pública própria; o Nginx "
        "recebe seu `hostport` pela propriedade `fromService` e encaminha `/api` e `/socket.io`. O "
        "MongoDB Atlas permanece externo e sua URI é solicitada no sincronismo do Blueprint."
    )
    add_body(
        document,
        "O Blueprint usa `sync: false` para valores fornecidos pelo operador e `generateValue: true` para "
        "segredos internos. `PUBLIC_APP_URL` e `CORS_ORIGINS` apontam para `https://notify-flow.onrender.com/`. "
        "Se o serviço for renomeado ou receber domínio próprio, ambos devem ser atualizados, bem como os "
        "callbacks dos provedores. O webhook WhatsApp usa "
        "`https://notify-flow.onrender.com/api/webhooks/whatsapp-cloud`; o Telegram deriva sua rota pública "
        "da mesma URL."
    )
    add_table(
        document,
        ["Domínio", "Variáveis exemplares", "Regra de produção"],
        [
            ("Aplicação", "NODE_ENV, PORT, API_PREFIX, PUBLIC_APP_URL", "URLs HTTPS e prefixos coerentes com o proxy."),
            ("Dados", "MONGODB_URI, REDIS_URL, REDIS_REQUIRED", "Atlas e Key Value privados; nunca versionar URI real."),
            ("Autenticação", "JWT_*_SECRET, PROFILE_JWT_SECRET, TTLs", "Segredos independentes e fortes; senha admin ≥ 12."),
            ("Proteção", "ENCRYPTION_KEY, SEARCH_HASH_KEY, RATE_LIMIT_*", "Chaves distintas, proxy confiável e limites ativos."),
            ("Canais", "WHATSAPP_*_VERSION, comandos, webhook secret", "Credenciais preferencialmente pela UI criptografada."),
        ],
        widths=[1.0, 2.3, 2.5],
        font_size=8.7,
    )
    add_caption(document, "Quadro 5 — Variáveis de ambiente por domínio")
    add_body(
        document,
        "A rota pública `/healthz` prova que o Nginx está de pé, enquanto `/api/health` atravessa o proxy "
        "e verifica API, Mongo e Redis. O Blueprint atual não declara health check próprio para o private "
        "service; portanto, recomenda-se monitorar externamente `/api/health` e adicionar uma verificação "
        "equivalente à API em uma revisão futura."
    )

    add_subheading(document, "4.11 Resultados dos testes", 2)
    add_body(
        document,
        "A bateria executada em 31 de julho de 2026 produziu os resultados da Tabela 1. Os totais foram "
        "obtidos pela execução das suítes, e não apenas por contagem estática dos arquivos."
    )
    add_table(
        document,
        ["Evidência", "Escopo", "Resultado observado"],
        [
            ("Testes da API", "34 arquivos: regras, filas, webhooks, segurança e persistência.", "305 aprovados; 0 falhas; 0 ignorados."),
            ("Testes do frontend", "31 arquivos: auth, serviços, páginas, diálogos e realtime.", "169 aprovados; 0 falhas."),
            ("Total automatizado", "API + frontend.", "474 aprovados; 0 falhas."),
            ("Lint da API", "ESLint sobre o código da API.", "Aprovado."),
            ("Build do frontend", "Vite em modo de produção.", "Aprovado; 376 módulos transformados."),
            ("Compose", "`docker compose config --quiet`.", "Configuração estática válida."),
        ],
        widths=[1.2, 2.8, 1.8],
        font_size=9.0,
    )
    add_caption(document, "Tabela 1 — Evidências de testes e validações")
    add_body(
        document,
        "Os 474 testes dão confiança sobre regressões de regra e contrato. A cobertura temática é "
        "particularmente relevante porque o núcleo do produto não é uma única tela, mas a combinação "
        "entre identidade, consentimento, fila e retorno do provedor. Casos testados incluem assinatura "
        "do webhook WhatsApp, deduplicação Telegram, renovação de sessão, grupos por convite, templates, "
        "permissões e campanhas com falhas parciais."
    )
    add_body(
        document,
        "O resultado não significa ausência de defeitos. As dependências externas são simuladas na maior "
        "parte da suíte, e a configuração local validada estava parada no instante da auditoria, com "
        "containers encerrados de forma limpa. Assim, o critério do Gate é atendido pela versão funcional "
        "já implantada e pelas evidências automatizadas, mas um plano de produção deve acrescentar smoke "
        "tests periódicos contra o ambiente real, monitoramento sintético e uma pipeline de integração "
        "contínua que bloqueie deploys quando teste ou build falhar."
    )

    add_subheading(document, "4.12 Discussão dos resultados e impacto", 2)
    add_body(
        document,
        "O resultado técnico mais importante é o isolamento de falhas. Em uma campanha com três canais, "
        "um contato pode não ter Telegram, outro pode ter e-mail revogado e um modelo WhatsApp pode ser "
        "rejeitado. O sistema não converte essas situações em uma falha única. Cada entrega possui estado, "
        "motivo e tentativas próprios. Essa propriedade melhora a continuidade operacional e permite "
        "corrigir somente o destino afetado."
    )
    add_body(
        document,
        "O segundo impacto é regulatório e de confiança. O fluxo privilegia uma ação do usuário: iniciar "
        "o bot, enviar o comando, compartilhar o próprio contato, validar e-mail ou usar a área pública. "
        "O administrador ainda pode ajustar permissões, mas a decisão fica registrada e a remoção exige "
        "confirmação. Dessa forma, o contato deixa de ser apenas um número em uma lista e passa a ser um "
        "titular com histórico, contexto de convite e controle de canal."
    )
    add_body(
        document,
        "O terceiro impacto é a redução do risco específico do WhatsApp. A solução não promete imunidade "
        "a bloqueios: modelos de baixa qualidade, denúncias ou políticas violadas continuam sujeitos à "
        "ação da Meta. O que muda é a base técnica. A comunicação iniciada pela organização passa por "
        "modelos oficiais, identidade remetente, limites, cobrança e recibos, em vez de depender de uma "
        "sessão automatizada do WhatsApp Web."
    )
    add_table(
        document,
        ["Risco ou limitação", "Controle atual", "Próxima ação"],
        [
            ("Indisponibilidade de provedor", "Fila, retries e erro por entrega.", "Circuit breaker e métricas de SLA."),
            ("Duplicidade em resposta ambígua", "Idempotência, claim e deduplicação.", "Chaves idempotentes provider-aware e reconciliação ampliada."),
            ("Escala horizontal", "Uma instância com worker e Socket.IO.", "Worker separado e adapter Redis do Socket.IO."),
            ("Backups no mesmo Mongo", "Snapshots criptografados em GridFS.", "PITR/backup externo e testes de restauração."),
            ("Crescimento de eventos", "TTLs em logs e recibos selecionados.", "Política de retenção para webhook events e campanhas."),
            ("Deploy sem CI", "Testes e build executáveis localmente.", "GitHub Actions antes do auto deploy."),
            ("Usabilidade não medida", "Fluxos responsivos e testes funcionais.", "Estudo com usuários e métricas de tarefa."),
            ("Segredo revelável ao admin", "Rota autenticada, limitada, sem cache e auditada.", "Acesso por função e rotação periódica."),
        ],
        widths=[1.55, 2.1, 2.15],
        font_size=8.5,
    )
    add_caption(document, "Tabela 2 — Riscos, controles e limitações")

    add_subheading(document, "4.13 Versão funcional demonstrável", 2)
    add_body(
        document,
        "A versão funcional está publicada em `https://notify-flow.onrender.com/`. A demonstração "
        "recomendada começa pelo painel administrativo: conferir a saúde, abrir configurações mascaradas, "
        "criar um convite, autorizar um contato, preparar templates e revisar uma campanha. Em seguida, "
        "o preceptor pode observar o trabalho na fila, abrir o detalhamento por contato e conferir no "
        "Meu Perfil o mesmo histórico sob a perspectiva do titular."
    )
    add_callout(
        document,
        "Roteiro de demonstração de 10 minutos",
        "1) mostrar configuração independente dos três canais; 2) abrir um convite e consentir; "
        "3) confirmar o contato e suas identidades; 4) selecionar um conjunto multicanal; 5) revisar e "
        "enfileirar; 6) acompanhar status e falha isolada; 7) abrir Meu Perfil e revogar um canal.",
        fill="EAF7FE",
        border="2AABEE",
    )


def add_conclusion(document: Document) -> None:
    add_heading(document, "5 CONSIDERAÇÕES FINAIS", 1)
    add_body(
        document,
        "O Gate 4 demonstrou que o Notify Flow atingiu o estágio de MVP funcional. A solução integra "
        "três canais, oferece interfaces administrativas e públicas, processa campanhas em fila, recebe "
        "eventos por webhook e atualiza o painel em tempo real. As 474 verificações automatizadas "
        "aprovadas, o lint, o build e a validação do Compose sustentam a conclusão técnica."
    )
    add_body(
        document,
        "A contribuição central é organizar uma comunicação que começa no consentimento e termina na "
        "auditoria. Para WhatsApp, isso significa trabalhar com a Cloud API, a janela de atendimento e "
        "os modelos aprovados; para Telegram, reconhecer a identidade observada pelo bot; para e-mail, "
        "validar o endereço. A fila impede que uma falha parcial apague o resultado da campanha e os "
        "recibos deixam clara a diferença entre envio, aceitação, entrega e leitura."
    )
    add_body(
        document,
        "Os próximos passos são submeter os fluxos a usuários externos, medir taxa de conclusão, tempo "
        "de tarefa e compreensão de consentimento; adicionar CI; separar o worker da API; adotar adapter "
        "Redis do Socket.IO; fortalecer retenção e backup; avaliar OAuth para Gmail; e implementar "
        "monitoramento sintético do ambiente Render. Uma distribuição desktop com Electron permanece "
        "como evolução posterior e não integra o entregável deste Gate."
    )
    add_body(
        document,
        "Conclui-se que o MVP é adequado para demonstração preliminar ao preceptor e oferece uma base "
        "coerente para evolução. Seu impacto esperado é reduzir o uso de integrações não oficiais, "
        "aumentar a rastreabilidade operacional e devolver ao contato controle explícito sobre os canais "
        "pelos quais aceita ser notificado."
    )


def add_references(document: Document) -> None:
    add_heading(document, "REFERÊNCIAS", 1)
    refs = [
        "BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD). Brasília, DF, 2018.",
        "BULLMQ. Retrying failing jobs. Disponível em: https://docs.bullmq.io/guide/retrying-failing-jobs. Acesso em: 31 jul. 2026.",
        "BULLMQ. Workers — concurrency. Disponível em: https://docs.bullmq.io/guide/workers/concurrency. Acesso em: 31 jul. 2026.",
        "DOCKER. Docker Compose documentation. Disponível em: https://docs.docker.com/compose/. Acesso em: 31 jul. 2026.",
        "GOOGLE. Sign in with app passwords. Disponível em: https://support.google.com/accounts/answer/185833. Acesso em: 31 jul. 2026.",
        "GOOGLE. Implement server-side authorization — Gmail API. Disponível em: https://developers.google.com/workspace/gmail/api/auth/web-server. Acesso em: 31 jul. 2026.",
        "META. WhatsApp Business Platform pricing. Disponível em: https://whatsappbusiness.com/products/platform-pricing/. Acesso em: 31 jul. 2026.",
        "META. WhatsApp Business Platform onboarding guide. Disponível em: https://whatsappbusiness.com/resources/resource-library/api-onboarding/. Acesso em: 31 jul. 2026.",
        "META. Documents for business verification. Disponível em: https://www.facebook.com/help/243868559497297/. Acesso em: 31 jul. 2026.",
        "RENDER. Blueprint YAML reference. Disponível em: https://render.com/docs/blueprint-spec. Acesso em: 31 jul. 2026.",
        "RENDER. Private services. Disponível em: https://render.com/docs/private-services. Acesso em: 31 jul. 2026.",
        "RENDER. Key Value. Disponível em: https://render.com/docs/key-value. Acesso em: 31 jul. 2026.",
        "TELEGRAM. Telegram Bot API. Disponível em: https://core.telegram.org/bots/api. Acesso em: 31 jul. 2026.",
    ]
    for ref in refs:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ref)
        set_run_font(r, 10.5)


def add_appendices(document: Document) -> None:
    add_heading(document, "APÊNDICE A — HISTÓRIAS DE USUÁRIO", 1)
    admin_rows = [
        ("A01", "Como administrador, quero configurar cada canal separadamente para implantar o sistema por etapas.", "Um canal incompleto não bloqueia os demais."),
        ("A02", "Como administrador, quero identificar novas interações em tempo real.", "Evento aparece sem recarregar e sem revelar segredo."),
        ("A03", "Como administrador, quero cadastrar templates por formulário.", "Não é necessário editar JSON."),
        ("A04", "Como administrador, quero agrupar templates por campanha.", "Conjunto possui até um template por canal."),
        ("A05", "Como administrador, quero selecionar contatos e grupos.", "Somente destinos identificados e permitidos são enviados."),
        ("A06", "Como administrador, quero revisar cada mensagem antes de enviar.", "Preview responsivo precede a confirmação."),
        ("A07", "Como administrador, quero acompanhar cada contato e canal.", "Status, tentativas e erro ficam detalhados."),
        ("A08", "Como administrador, quero sincronizar grupos por convite.", "Novos contatos autorizados são incluídos sem duplicidade."),
        ("A09", "Como administrador, quero responder no WhatsApp.", "Texto livre só é permitido na janela de 24 horas."),
        ("A10", "Como administrador, quero versionar termos e privacidade.", "Documento vigente é apresentado no convite."),
        ("A11", "Como administrador, quero auditar links de acesso.", "Uso, validade e revogação são registrados."),
        ("A12", "Como administrador, quero que uma falha não interrompa a campanha.", "Cada entrega é processada e registrada isoladamente."),
    ]
    add_table(
        document,
        ["ID", "História", "Critério resumido"],
        admin_rows,
        widths=[0.45, 3.35, 2.0],
        font_size=8.3,
    )
    add_caption(document, "Tabela 3 — Histórias de usuário do administrador")

    contact_rows = [
        ("U01", "Como contato, quero escolher quais canais autorizo.", "Permissões independentes e revogáveis."),
        ("U02", "Como contato, quero iniciar a autorização pelo canal.", "Comando ou convite resulta em identidade observada."),
        ("U03", "Como contato, quero vincular Telegram e WhatsApp com segurança.", "Telefone só é aceito quando compartilhado pelo próprio remetente."),
        ("U04", "Como contato, quero validar meu e-mail no chat.", "Código expira em 15 min e reenvio respeita 2 min."),
        ("U05", "Como contato, quero entrar no Meu Perfil sem senha permanente.", "Link assinado, temporário e de uso único."),
        ("U06", "Como contato, quero consultar minhas entregas.", "Histórico é restrito ao próprio contato."),
        ("U07", "Como contato, quero compreender convites e grupos.", "Perfil lista vínculos e participantes mascarados."),
        ("U08", "Como contato, quero remover meu vínculo.", "Ação afeta somente o próprio cadastro e exige confirmação."),
        ("U09", "Como contato, quero revogar um canal.", "Novos envios são bloqueados e a decisão fica auditada."),
        ("U10", "Como visitante, quero uma página de convite responsiva.", "Links e QR funcionam em desktop e celular."),
    ]
    add_table(
        document,
        ["ID", "História", "Critério resumido"],
        contact_rows,
        widths=[0.45, 3.35, 2.0],
        font_size=8.3,
    )
    add_caption(document, "Tabela 4 — Histórias de usuário do contato")

    add_heading(document, "APÊNDICE B — CHECKLIST DE CONFIGURAÇÃO", 1)
    add_subheading(document, "B.1 WhatsApp Cloud", 2)
    for item in [
        "Confirmar empresa, Página, portfólio e verificações de segurança.",
        "Registrar número elegível e associar forma de pagamento.",
        "Criar Meta App, adicionar o produto WhatsApp e selecionar a WABA.",
        "Obter Phone Number ID, Business Account ID e token apropriado.",
        "Cadastrar App Secret e verify token no Notify Flow.",
        "Configurar callback HTTPS `/api/webhooks/whatsapp-cloud` e assinar `messages`.",
        "Criar modelos no WhatsApp Manager e aguardar aprovação.",
        "Cadastrar no Notify Flow o mesmo nome, idioma e parâmetros.",
        "Testar número de teste e, depois, número de produção com destinatário autorizado.",
    ]:
        add_bullet(document, item)
    add_subheading(document, "B.2 Telegram", 2)
    for item in [
        "Criar o bot no @BotFather e armazenar o token como segredo.",
        "Salvar o token; conferir nome e username identificados por getMe.",
        "Garantir URL pública HTTPS e registrar webhook com secret token.",
        "Iniciar o bot, testar /notify-me, /verify-me e /login.",
        "Compartilhar o próprio contato para validar o vínculo de telefone.",
        "Testar texto, foto/vídeo por URL segura e menu de botões.",
    ]:
        add_bullet(document, item)
    add_subheading(document, "B.3 Gmail", 2)
    for item in [
        "Ativar verificação em duas etapas na Conta Google.",
        "Gerar senha de aplicativo quando o recurso estiver disponível.",
        "Cadastrar conta Gmail, senha de aplicativo e remetente autorizado.",
        "Enviar teste individual e campanha com endereço autorizado.",
        "Confirmar que falha SMTP de um contato não interrompe os demais.",
    ]:
        add_bullet(document, item)
    add_subheading(document, "B.4 Produção", 2)
    for item in [
        "Criar MongoDB Atlas com usuário de menor privilégio, rede e backup/PITR.",
        "Sincronizar o Blueprint e fornecer MONGODB_URI e credenciais administrativas.",
        "Verificar que frontend, API privada e Key Value estão na mesma região.",
        "Testar `/healthz` e `/api/health`.",
        "Atualizar callbacks ao alterar domínio.",
        "Executar smoke test de convite, consentimento, campanha e Meu Perfil.",
        "Monitorar memória do Key Value, fila, erros por provedor e qualidade da WABA.",
    ]:
        add_bullet(document, item)

    add_heading(document, "APÊNDICE C — MATRIZ DE EVIDÊNCIAS", 1)
    add_table(
        document,
        ["Área", "Evidência no projeto", "Validação do Gate 4"],
        [
            ("API", "`api/src`, 21 modelos, 18 módulos de rota, 34 arquivos de teste.", "305 testes e lint aprovados."),
            ("Frontend", "14 rotas funcionais, 31 arquivos de teste, build Quasar/Vite.", "169 testes e build aprovados."),
            ("Fila", "BullMQ, Redis obrigatório, Mongo durável e recovery sweep.", "Retries e falha parcial cobertos por testes."),
            ("Webhooks", "HMAC WhatsApp, secret Telegram e deduplicação.", "Contratos e validações testados."),
            ("Infra local", "`docker-compose.yml` com quatro serviços.", "Compose validado estaticamente."),
            ("Infra produção", "`render.yaml`: web, private service e Key Value.", "Blueprint inspecionado e deploy disponível."),
            ("Governança", "ConsentEvent, termos, Meu Perfil e pseudonimização.", "Fluxos e políticas cobertos por testes."),
        ],
        widths=[1.0, 2.9, 1.9],
        font_size=8.8,
    )
    add_caption(document, "Tabela 5 — Matriz de evidências do entregável")


def normalize_document(document: Document) -> None:
    # Keep page geometry from the reference but guarantee expected margins.
    for index, section in enumerate(document.sections):
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        section.different_first_page_header_footer = index == 0

    # Accessibility: every retained/generated inline image receives alt text.
    for idx, shape in enumerate(document.inline_shapes, start=1):
        doc_pr = shape._inline.docPr
        if not doc_pr.get("descr"):
            doc_pr.set("descr", f"Imagem institucional ou figura técnica {idx} do relatório Notify Flow")


def build() -> Path:
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    if hashlib.sha256(REFERENCE.read_bytes()).hexdigest().upper() != (
        "70103C81C778FE548FDED375E82327C8E7A82108BA86387A551517B18C53F306"
    ):
        raise RuntimeError("O template de referência foi alterado desde a inspeção.")

    diagrams = save_diagrams()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    working = WORK / "working-report.docx"
    shutil.copy2(REFERENCE, working)
    document = Document(working)
    configure_styles(document)
    cover_and_title_page(document)
    remove_after_paragraph(document, 61)
    add_front_matter(document)

    # Main text starts in a new Word section so the front matter remains distinct.
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True

    add_introduction(document)
    add_theory(document)
    add_methodology(document)
    add_development_results(document, diagrams)
    add_conclusion(document)
    add_references(document)
    add_appendices(document)

    patch_header_title(document)
    normalize_document(document)
    set_update_fields(document)
    set_core_properties(document)
    document.save(OUTPUT)
    print(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    build()
