from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCX = Path(
    r"C:\Users\samuel\tcc\notify-app\docs"
    r"\Relatorio-de-Impacto-Notify-Flow-Gate-4-Samuel-Victor-Oliveira-Lima.docx"
)
OUT = Path(
    r"C:\Users\samuel\tcc\notify-app\.codex-tmp"
    r"\notify-flow-impact-report\qa-summary.json"
)


def main() -> int:
    failures: list[str] = []
    warnings: list[str] = []

    if not DOCX.exists() or DOCX.stat().st_size < 100_000:
        failures.append("Arquivo ausente ou pequeno demais para conter o relatório e as figuras.")
        report = {"failures": failures}
        OUT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        return 1

    document = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    normalized = re.sub(r"\s+", " ", text)
    words = re.findall(r"\b[\wÀ-ÿ][\wÀ-ÿ0-9_/-]*\b", text)
    headings = [
        (paragraph.style.name, paragraph.text.strip())
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading") and paragraph.text.strip()
    ]

    required = [
        "Samuel Victor Oliveira Lima",
        "Leonardo Oliveira",
        "RESUMO EXECUTIVO",
        "4 DESENVOLVIMENTO, RESULTADOS E DISCUSSÃO",
        "474 testes",
        "305 testes",
        "169",
        "https://notify-flow.onrender.com/",
        "WhatsApp Cloud",
        "Telegram",
        "Gmail",
        "BullMQ",
        "Redis",
        "Docker Compose",
        "Render",
        "Meu Perfil",
        "jaspers_market_plain_text_v1",
        "jaspers_market_order_confirmation_v1",
        "3p_direct_integration_test_template",
        "APÊNDICE A",
    ]
    for value in required:
        if value not in text:
            failures.append(f"Conteúdo obrigatório ausente: {value}")

    forbidden = [
        "NOME DO ESTUDANTE",
        "TÍTULO DO TRABALHO DE CONCLUSÃO",
        "Texto texto texto",
        "Listar as figuras contidas",
        "Colocar o Título",
        "nome do orientador",
        "EAAY",
        "EAAO",
        "mongodb+srv://",
        "wrm78",
        "Guitas@",
    ]
    for value in forbidden:
        if value.lower() in text.lower():
            failures.append(f"Placeholder ou segredo aparente encontrado no texto: {value}")

    if len(words) < 5_500:
        failures.append(f"Relatório curto demais para o escopo solicitado: {len(words)} palavras.")
    if len(headings) < 35:
        failures.append(f"Hierarquia de títulos insuficiente: {len(headings)}.")
    if len(document.tables) < 12:
        failures.append(f"Quantidade inesperada de tabelas/quadros: {len(document.tables)}.")
    if len(document.inline_shapes) != 6:
        failures.append(f"Quantidade inesperada de imagens: {len(document.inline_shapes)}.")
    if len(document.sections) != 2:
        failures.append(f"Quantidade inesperada de seções Word: {len(document.sections)}.")

    with zipfile.ZipFile(DOCX) as archive:
        names = set(archive.namelist())
        xml_parts = [
            name
            for name in names
            if name.endswith(".xml") or name.endswith(".rels")
        ]
        xml_text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore") for name in xml_parts
        )

        for value in forbidden:
            if value.lower() in xml_text.lower():
                failures.append(f"Placeholder ou segredo aparente encontrado no pacote: {value}")

        if any(name.startswith("word/comments") for name in names):
            failures.append("Partes de comentários ainda existem no pacote.")
        if re.search(r"<w:(?:ins|del)(?:\s|>)", xml_text):
            failures.append("Marcas de revisão ainda existem no documento.")
        if "TOC \\o" not in xml_text:
            failures.append("Campo TOC não encontrado.")
        if not re.search(r">\s*PAGE\s*<", xml_text):
            failures.append("Campo PAGE não encontrado.")
        if "w:updateFields" not in xml_text:
            failures.append("Documento não solicita atualização de campos.")

        media = [name for name in names if name.startswith("word/media/")]
        if len(media) != 6:
            failures.append(f"Quantidade inesperada de mídias no pacote: {len(media)}.")

        doc_pr_count = len(re.findall(r"<wp:docPr\b", xml_text))
        alt_count = len(re.findall(r'<wp:docPr\b[^>]*\bdescr="[^"]+"', xml_text))
        if doc_pr_count != alt_count:
            failures.append(
                f"Nem todas as imagens possuem descrição alternativa: {alt_count}/{doc_pr_count}."
            )

        leaked_patterns = [
            r"(?i)authorization:\s*bearer\s+[A-Za-z0-9._-]{20,}",
            r"(?i)redis://[^<\s]+",
            r"(?i)mongodb(?:\+srv)?://[^<\s]+",
            r"(?i)jwt_[a-z_]*secret\s*=\s*[^<\s]+",
        ]
        for pattern in leaked_patterns:
            if re.search(pattern, xml_text):
                failures.append(f"Padrão de credencial aparente encontrado: {pattern}")

    heading_texts = [item[1] for item in headings]
    if len(heading_texts) != len(set(heading_texts)):
        duplicates = sorted({h for h in heading_texts if heading_texts.count(h) > 1})
        warnings.append(f"Títulos repetidos: {duplicates}")

    figure_captions = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Figura ")
        and " — " in paragraph.text
    ]
    expected_figures = [f"Figura {number} —" for number in range(1, 6)]
    for prefix in expected_figures:
        if not any(caption.startswith(prefix) for caption in figure_captions):
            failures.append(f"Legenda ausente: {prefix}")

    report = {
        "file": str(DOCX),
        "size_bytes": DOCX.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "words": len(words),
        "headings": len(headings),
        "tables": len(document.tables),
        "inline_images": len(document.inline_shapes),
        "sections": len(document.sections),
        "figure_captions": figure_captions,
        "failures": failures,
        "warnings": warnings,
    }
    OUT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
